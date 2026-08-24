"""Docling layout parsing and Ollama-backed OCR specialist."""

from __future__ import annotations

import base64
import json
import os
import re
from io import BytesIO
from pathlib import Path
from uuid import uuid4

import requests
from docx import Document as WordDocument

from src.custodian_backend import CustodianClient, CustodianError
from src.document_attachments import OCR_UPLOAD_DIR
from src.workspace_policy import canonical_workspace

DEFAULT_SURYA_MODEL = "hf.co/datalab-to/surya-ocr-2-gguf:latest"
DEFAULT_GLM_MODEL = "glm-ocr:bf16"
MAX_OCR_BYTES = 25 * 1024 * 1024


def approved_document_path(document_ref: str, workspace: str | None) -> Path:
    """Resolve only a selected workspace file or an approved upload reference."""
    if document_ref.startswith("upload:"):
        name = document_ref.removeprefix("upload:")
        if not name or Path(name).name != name:
            raise ValueError("Invalid OCR upload reference")
        path = (OCR_UPLOAD_DIR / name).resolve()
        root = OCR_UPLOAD_DIR
    else:
        if not workspace:
            raise ValueError("An OCR document path requires a selected workspace")
        root = canonical_workspace(workspace)
        candidate = Path(document_ref).expanduser()
        path = (candidate if candidate.is_absolute() else root / candidate).resolve()
    try:
        path.relative_to(root)
    except ValueError as exc:
        raise ValueError("OCR document is outside the approved directory") from exc
    if not path.is_file():
        raise ValueError("OCR document was not found")
    if path.stat().st_size > MAX_OCR_BYTES:
        raise ValueError("OCR document exceeds the 25 MB limit")
    return path


def _workspace_document_reference(
    document_ref: str, workspace: str
) -> tuple[Path, str]:
    root = canonical_workspace(workspace)
    candidate = Path(document_ref).expanduser()
    if candidate.is_absolute():
        normalized = Path(os.path.normpath(str(candidate)))
        if str(normalized) != str(candidate):
            raise ValueError("OCR document path must be canonical")
        try:
            relative = normalized.relative_to(root)
        except ValueError:
            relative = Path(str(normalized).lstrip("/"))
    else:
        if ".." in candidate.parts:
            raise ValueError("OCR document is outside the selected workspace")
        relative = candidate
    if not relative.parts:
        raise ValueError("An OCR document file is required")
    return root, f"/{relative.as_posix().lstrip('/')}"


def _stage_workspace_document(document_ref: str, workspace: str) -> str:
    root, virtual_path = _workspace_document_reference(document_ref, workspace)
    try:
        result = CustodianClient(str(root), timeout=60).action(
            "stage_ocr_document", path=virtual_path
        )
    except CustodianError as exc:
        raise ValueError("OCR document staging failed") from exc
    if result.get("ok") is not True:
        raise ValueError(str(result.get("error") or "OCR document staging failed"))
    reference = str(result.get("reference") or "")
    name = reference.removeprefix("upload:")
    if not reference.startswith("upload:") or not name or Path(name).name != name:
        raise ValueError("Custodian returned an invalid OCR upload reference")
    return reference


def _write_workspace_output(
    document_ref: str,
    workspace: str,
    content: str | bytes,
    output_format: str,
) -> str:
    root, virtual_path = _workspace_document_reference(document_ref, workspace)
    payload = {"path": virtual_path, "output_format": output_format}
    if isinstance(content, bytes):
        if output_format != "docx":
            raise ValueError("Binary OCR output requires the docx format")
        payload["content_base64"] = base64.b64encode(content).decode("ascii")
    else:
        payload["content"] = content
    try:
        result = CustodianClient(str(root), timeout=60).action(
            "write_ocr_output", **payload
        )
    except CustodianError as exc:
        raise ValueError("OCR workspace output failed") from exc
    if result.get("ok") is not True:
        raise ValueError(str(result.get("error") or "OCR workspace output failed"))
    output_path = str(result.get("path") or "")
    if not output_path.startswith("/") or ".." in Path(output_path).parts:
        raise ValueError("Custodian returned an invalid OCR output path")
    return output_path


def resolve_document_path(document_ref: str, workspace: str | None) -> Path:
    try:
        return approved_document_path(document_ref, workspace)
    except ValueError:
        if document_ref.startswith("upload:") or not workspace:
            raise
    staged_reference = _stage_workspace_document(document_ref, workspace)
    return approved_document_path(staged_reference, workspace)


def _docling_layout(path: Path, render_dir: Path) -> tuple[str, list[Path]]:
    try:
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.document_converter import (
            DocumentConverter,
            PdfFormatOption,
        )
    except ImportError as exc:
        raise RuntimeError("Docling is required in the runtime image for OCR") from exc

    pipeline_options = PdfPipelineOptions(
        do_ocr=True,
        do_table_structure=True,
        generate_page_images=True,
        images_scale=2.0,
    )
    converter = DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
        }
    )
    document = converter.convert(str(path)).document
    markdown = document.export_to_markdown().strip()
    if not markdown:
        raise RuntimeError("Docling did not produce an ordered document")

    pages: list[Path] = []
    render_dir.mkdir(parents=True, exist_ok=True)
    for number, page in sorted(getattr(document, "pages", {}).items()):
        page_image = getattr(page, "image", None)
        image = getattr(page_image, "pil_image", page_image)
        if image is not None:
            target = render_dir / f"page-{number}.png"
            image.save(target)
            pages.append(target)
    return markdown, pages


def _docx_bytes(markdown: str) -> bytes:
    document = WordDocument()
    for block in re.split(r"\n[ \t]*\n", markdown):
        if not block:
            continue
        heading = re.fullmatch(r"(#{1,6})[ \t]+(.+)", block)
        if heading:
            document.add_heading(heading.group(2), level=len(heading.group(1)))
            continue
        paragraph = document.add_paragraph()
        lines = block.splitlines()
        for index, line in enumerate(lines):
            paragraph.add_run(line)
            if index < len(lines) - 1:
                paragraph.add_run().add_break()
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _ollama_ocr(image: Path, model: str, base_url: str, timeout: int) -> str:
    encoded = base64.b64encode(image.read_bytes()).decode("ascii")
    response = requests.post(
        f"{base_url.rstrip('/')}/api/chat",
        json={
            "model": model,
            "messages": [
                {
                    "role": "user",
                    "content": "Transcribe this page faithfully.",
                    "images": [encoded],
                }
            ],
            "stream": False,
            "keep_alive": 0,
        },
        timeout=timeout,
    )
    response.raise_for_status()
    content = response.json().get("message", {}).get("content", "")
    if not isinstance(content, str):
        raise ValueError("OCR verifier returned non-text content")
    max_chars = int(os.getenv("OCR_MAX_VERIFICATION_CHARS", "200000"))
    if len(content) > max_chars:
        raise ValueError("OCR verifier output exceeded its configured limit")
    return content


def _run_verification_phase(
    pages: list[Path],
    phase: str,
    model: str,
    base_url: str,
    timeout: int,
) -> tuple[list[str | None], list[dict]]:
    outputs: list[str | None] = []
    failures: list[dict] = []
    for page_number, page in enumerate(pages, 1):
        try:
            outputs.append(_ollama_ocr(page, model, base_url, timeout))
        except Exception as exc:
            outputs.append(None)
            failures.append(
                {
                    "page": page_number,
                    "phase": phase,
                    "error": type(exc).__name__,
                }
            )
    return outputs, failures


def run_ocr(
    task: str,
    document_ref: str,
    workspace: str | None,
    output_format: str = "markdown",
) -> dict:
    if output_format not in {"markdown", "json", "structured"}:
        raise ValueError("output_format must be markdown, json, or structured")
    path = resolve_document_path(document_ref, workspace)
    run_dir = (
        Path(
            os.getenv("OCR_ARTIFACT_DIR", str(path.parent / ".ocr-artifacts"))
        ).resolve()
        / uuid4().hex
    )
    layout, pages = _docling_layout(path, run_dir / "pages")
    docx_content = _docx_bytes(layout)
    docx_artifact = run_dir / "result.docx"
    docx_artifact.parent.mkdir(parents=True, exist_ok=True)
    docx_artifact.write_bytes(docx_content)
    base_url = os.getenv("OLLAMA_BASE_URL", "http://127.0.0.1:11434")
    timeout = int(os.getenv("OCR_OLLAMA_TIMEOUT", "300"))
    surya_model = os.getenv("OCR_SURYA_MODEL", DEFAULT_SURYA_MODEL)
    glm_model = os.getenv("OCR_GLM_MODEL", DEFAULT_GLM_MODEL)
    surya, surya_failures = _run_verification_phase(
        pages, "surya", surya_model, base_url, timeout
    )
    glm, glm_failures = _run_verification_phase(
        pages, "glm_ocr", glm_model, base_url, timeout
    )
    verification_failures = [*surya_failures, *glm_failures]
    disagreements = [
        {"page": index}
        for index, (left, right) in enumerate(zip(surya, glm, strict=True), 1)
        if left is not None
        and right is not None
        and re.sub(r"\s+", " ", left).strip() != re.sub(r"\s+", " ", right).strip()
    ]
    models = {"surya": surya_model, "glm_ocr": glm_model}
    result = {
        "task": task,
        "document": str(path),
        "normalized": layout,
        "layout_markdown": layout,
        "layout_authority": "docling",
        "model_role": "text_verification_only",
        "models": models,
        "surya": surya,
        "glm_ocr": glm,
        "disagreements": disagreements,
        "model_disagreement": bool(disagreements),
        "verification_failures": verification_failures,
    }
    artifact = run_dir / (
        "result.json" if output_format in {"json", "structured"} else "result.md"
    )
    artifact_content = (
        json.dumps(result, indent=2) if artifact.suffix == ".json" else layout
    )
    artifact.parent.mkdir(parents=True, exist_ok=True)
    artifact.write_text(artifact_content, encoding="utf-8")

    workspace_content = layout
    if output_format in {"json", "structured"}:
        workspace_content = json.dumps(
            {
                "normalized": layout,
                "layout_markdown": layout,
                "layout_authority": "docling",
                "model_role": "text_verification_only",
                "page_count": len(pages),
                "models": models,
                "disagreements": disagreements,
                "model_disagreement": bool(disagreements),
                "verification_failures": verification_failures,
            },
            indent=2,
        )

    workspace_output = None
    workspace_docx_output = None
    if workspace and not document_ref.startswith("upload:"):
        workspace_output = _write_workspace_output(
            document_ref, workspace, workspace_content, output_format
        )
        workspace_docx_output = _write_workspace_output(
            document_ref, workspace, docx_content, "docx"
        )

    manifest = run_dir / "manifest.json"
    manifest.write_text(
        json.dumps(
            {
                "artifact": str(artifact),
                "docx_artifact": str(docx_artifact),
                "workspace_output": workspace_output,
                "workspace_docx_output": workspace_docx_output,
                "pages": [str(page) for page in pages],
                "layout_authority": "docling",
                "model_role": "text_verification_only",
                "models": {
                    "surya": surya_model,
                    "glm_ocr": glm_model,
                },
                "model_disagreement": bool(disagreements),
                "verification_failures": verification_failures,
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    result.update(
        {
            "artifact_path": str(artifact),
            "docx_artifact_path": str(docx_artifact),
            "manifest_path": str(manifest),
            "workspace_output": workspace_output,
            "workspace_docx_output": workspace_docx_output,
        }
    )
    return result


def specialist_message(result: dict, output_format: str) -> str:
    page_count = len(result.get("surya", []))
    failure_count = len(result.get("verification_failures", []))
    disagreement_count = len(result.get("disagreements", []))
    page_word = "page" if page_count == 1 else "pages"
    if failure_count:
        verification = (
            f"Model verification was incomplete for {failure_count} page checks."
        )
    elif disagreement_count:
        verification = (
            f"The verification models disagreed on {disagreement_count} "
            f"of {page_count} {page_word}."
        )
    elif page_count:
        verification = (
            f"The verification models agreed on all {page_count} {page_word}."
        )
    else:
        verification = "No rendered pages were available for model verification."
    output_path = result.get("workspace_output") or result["artifact_path"]
    docx_path = result.get("workspace_docx_output") or result["docx_artifact_path"]
    return (
        "OCR complete. Docling determined the layout and reading order. "
        f"Outputs: {output_path} and {docx_path}. {verification}"
    )
