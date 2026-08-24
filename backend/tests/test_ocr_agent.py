import base64
import json
from io import BytesIO
from pathlib import Path
from unittest.mock import Mock

import pytest
from docx import Document as WordDocument
from langchain_core.messages import AIMessage
from langgraph.types import Command

from src import ocr_agent
from src.chat_ui import create_chat_ui
from src.jasper_agent import transfer_to_ocr


def test_ocr_path_is_confined_to_workspace(tmp_path):
    document = tmp_path / "scan.pdf"
    document.write_bytes(b"pdf")
    assert ocr_agent.approved_document_path("scan.pdf", str(tmp_path)) == document
    with pytest.raises(ValueError):
        ocr_agent.approved_document_path("../scan.pdf", str(tmp_path))


def test_workspace_document_is_staged_through_custodian(monkeypatch):
    monkeypatch.setenv("WORKSPACE_AUTHORIZED_ROOTS", "/Volumes/Storage")
    calls = []

    class Client:
        def __init__(self, workspace, timeout):
            calls.append((workspace, timeout))

        def action(self, action, **payload):
            calls.append((action, payload))
            return {"ok": True, "reference": "upload:staged-RUBTTI.pdf"}

    monkeypatch.setattr(ocr_agent, "CustodianClient", Client)

    reference = ocr_agent._stage_workspace_document(
        "/Volumes/Storage/example/RUBTTI.pdf", "/Volumes/Storage/example"
    )

    assert reference == "upload:staged-RUBTTI.pdf"
    assert calls == [
        ("/Volumes/Storage/example", 60),
        ("stage_ocr_document", {"path": "/RUBTTI.pdf"}),
    ]


def test_workspace_output_is_written_through_custodian(monkeypatch):
    monkeypatch.setenv("WORKSPACE_AUTHORIZED_ROOTS", "/Volumes/Storage")
    calls = []

    class Client:
        def __init__(self, workspace, timeout):
            calls.append((workspace, timeout))

        def action(self, action, **payload):
            calls.append((action, payload))
            return {"ok": True, "path": "/RUBTTI.ocr.md"}

    monkeypatch.setattr(ocr_agent, "CustodianClient", Client)

    output = ocr_agent._write_workspace_output(
        "/Volumes/Storage/example/RUBTTI.pdf",
        "/Volumes/Storage/example",
        "recognized text",
        "markdown",
    )

    assert output == "/RUBTTI.ocr.md"
    assert calls == [
        ("/Volumes/Storage/example", 60),
        (
            "write_ocr_output",
            {
                "path": "/RUBTTI.pdf",
                "content": "recognized text",
                "output_format": "markdown",
            },
        ),
    ]


def test_workspace_docx_is_base64_encoded_for_custodian(monkeypatch):
    monkeypatch.setenv("WORKSPACE_AUTHORIZED_ROOTS", "/Volumes/Storage")
    calls = []

    class Client:
        def __init__(self, workspace, timeout):
            pass

        def action(self, action, **payload):
            calls.append((action, payload))
            return {"ok": True, "path": "/RUBTTI.ocr.docx"}

    monkeypatch.setattr(ocr_agent, "CustodianClient", Client)
    content = ocr_agent._docx_bytes("Paragraph.")

    output = ocr_agent._write_workspace_output(
        "/Volumes/Storage/example/RUBTTI.pdf",
        "/Volumes/Storage/example",
        content,
        "docx",
    )

    assert output == "/RUBTTI.ocr.docx"
    assert calls[0][0] == "write_ocr_output"
    assert calls[0][1]["output_format"] == "docx"
    assert base64.b64decode(calls[0][1]["content_base64"]) == content
    assert "content" not in calls[0][1]


def test_document_resolution_falls_back_to_staged_upload(monkeypatch):
    staged_path = Path("/shared/staged-RUBTTI.pdf")
    calls = []

    def approved(reference, workspace):
        calls.append((reference, workspace))
        if not reference.startswith("upload:"):
            raise ValueError("not visible in the Agent Server")
        return staged_path

    monkeypatch.setattr(ocr_agent, "approved_document_path", approved)
    monkeypatch.setattr(
        ocr_agent,
        "_stage_workspace_document",
        lambda reference, workspace: "upload:staged-RUBTTI.pdf",
    )

    assert (
        ocr_agent.resolve_document_path(
            "/Volumes/Storage/example/RUBTTI.pdf", "/Volumes/Storage/example"
        )
        == staged_path
    )
    assert calls == [
        ("/Volumes/Storage/example/RUBTTI.pdf", "/Volumes/Storage/example"),
        ("upload:staged-RUBTTI.pdf", "/Volumes/Storage/example"),
    ]


def test_transfer_schema_and_command_route():
    schema = transfer_to_ocr.args_schema.model_json_schema()
    assert schema["properties"]["output_format"]["enum"] == [
        "markdown",
        "json",
        "structured",
    ]
    runtime = type(
        "Runtime",
        (),
        {
            "state": {"messages": [AIMessage(content="handoff")]},
            "tool_call_id": "call-1",
        },
    )()
    command = transfer_to_ocr.func(
        "read tables", "upload:x.pdf", "json", runtime=runtime
    )
    assert isinstance(command, Command)
    assert command.goto == "ocr"
    assert command.update["ocr_output_format"] == "json"


def test_ollama_ocr_requests_model_unload(monkeypatch, tmp_path):
    image = tmp_path / "page.png"
    image.write_bytes(b"image")
    response = Mock()
    response.json.return_value = {"message": {"content": "text"}}
    post = Mock(return_value=response)
    monkeypatch.setattr(ocr_agent.requests, "post", post)

    assert ocr_agent._ollama_ocr(image, "surya", "http://ollama", 12) == "text"
    payload = post.call_args.kwargs["json"]
    assert payload["keep_alive"] == 0
    assert post.call_args.kwargs["timeout"] == 12


@pytest.mark.parametrize("content", [None, ["not", "text"]])
def test_ollama_ocr_rejects_non_text_results(monkeypatch, tmp_path, content):
    image = tmp_path / "page.png"
    image.write_bytes(b"image")
    response = Mock()
    response.json.return_value = {"message": {"content": content}}
    monkeypatch.setattr(ocr_agent.requests, "post", Mock(return_value=response))

    with pytest.raises(ValueError, match="non-text"):
        ocr_agent._ollama_ocr(image, "surya", "http://ollama", 12)


def test_docx_preserves_docling_headings_and_paragraph_boundaries():
    content = ocr_agent._docx_bytes(
        "# Heading\n\nFirst line\ncontinues.\n\nSecond paragraph."
    )
    document = WordDocument(BytesIO(content))

    assert [paragraph.text for paragraph in document.paragraphs] == [
        "Heading",
        "First line\ncontinues.",
        "Second paragraph.",
    ]
    assert document.paragraphs[0].style.name == "Heading 1"


def test_ocr_runs_surya_phase_before_glm_verification(monkeypatch, tmp_path):
    document = tmp_path / "scan.pdf"
    document.write_bytes(b"pdf")
    pages = [tmp_path / "page-1.png", tmp_path / "page-2.png"]
    monkeypatch.setattr(
        ocr_agent, "_docling_layout", lambda path, render: ("layout", pages)
    )
    calls = []

    def mocked_ocr(image, model, base_url, timeout):
        calls.append(model)
        return f"{model}:{image.name}"

    monkeypatch.setattr(ocr_agent, "_ollama_ocr", mocked_ocr)
    workspace_writes = []

    def mocked_workspace_write(*args):
        workspace_writes.append(args)
        return "/scan.ocr.docx" if args[3] == "docx" else "/scan.ocr.md"

    monkeypatch.setattr(ocr_agent, "_write_workspace_output", mocked_workspace_write)
    result = ocr_agent.run_ocr("read", str(document), str(tmp_path))

    assert calls == [
        ocr_agent.DEFAULT_SURYA_MODEL,
        ocr_agent.DEFAULT_SURYA_MODEL,
        ocr_agent.DEFAULT_GLM_MODEL,
        ocr_agent.DEFAULT_GLM_MODEL,
    ]
    assert result["surya"] == [
        f"{ocr_agent.DEFAULT_SURYA_MODEL}:page-1.png",
        f"{ocr_agent.DEFAULT_SURYA_MODEL}:page-2.png",
    ]
    assert result["glm_ocr"] == [
        f"{ocr_agent.DEFAULT_GLM_MODEL}:page-1.png",
        f"{ocr_agent.DEFAULT_GLM_MODEL}:page-2.png",
    ]
    assert len(result["disagreements"]) == 2
    assert result["normalized"] == "layout"
    assert result["layout_authority"] == "docling"
    assert result["model_role"] == "text_verification_only"
    assert result["workspace_output"] == "/scan.ocr.md"
    assert result["workspace_docx_output"] == "/scan.ocr.docx"
    assert Path(result["docx_artifact_path"]).is_file()
    assert workspace_writes[0][0:2] == (str(document), str(tmp_path))
    assert workspace_writes[0][2] == "layout"
    assert workspace_writes[0][3] == "markdown"
    assert workspace_writes[1][0:2] == (str(document), str(tmp_path))
    assert isinstance(workspace_writes[1][2], bytes)
    assert workspace_writes[1][3] == "docx"
    assert WordDocument(BytesIO(workspace_writes[1][2])).paragraphs[0].text == "layout"


def test_model_failure_cannot_replace_or_block_docling_layout(monkeypatch, tmp_path):
    document = tmp_path / "scan.pdf"
    document.write_bytes(b"pdf")
    layout = "First paragraph.\n\nSecond paragraph."
    monkeypatch.setattr(
        ocr_agent, "_docling_layout", lambda path, render: (layout, [path])
    )

    def mocked_ocr(image, model, base_url, timeout):
        if model == ocr_agent.DEFAULT_SURYA_MODEL:
            raise TimeoutError("verification unavailable")
        return "model text"

    monkeypatch.setattr(ocr_agent, "_ollama_ocr", mocked_ocr)
    monkeypatch.setattr(
        ocr_agent,
        "_write_workspace_output",
        lambda *args: "/scan.ocr.md",
    )

    result = ocr_agent.run_ocr("read", str(document), str(tmp_path))

    assert result["normalized"] == layout
    assert result["surya"] == [None]
    assert result["glm_ocr"] == ["model text"]
    assert result["disagreements"] == []
    assert result["verification_failures"] == [
        {"page": 1, "phase": "surya", "error": "TimeoutError"}
    ]
    assert "verification was incomplete" in ocr_agent.specialist_message(
        result, "markdown"
    )


def test_json_output_keeps_docling_layout_and_chat_gets_only_summary(
    monkeypatch, tmp_path
):
    document = tmp_path / "scan.pdf"
    document.write_bytes(b"pdf")
    artifact_dir = tmp_path / "artifacts"
    layout = "# Heading\n\nParagraph preserved."
    monkeypatch.setenv("OCR_ARTIFACT_DIR", str(artifact_dir))
    monkeypatch.setattr(
        ocr_agent, "_docling_layout", lambda path, render: (layout, [path])
    )
    monkeypatch.setattr(
        ocr_agent, "_ollama_ocr", lambda image, model, base, timeout: "text"
    )
    workspace_writes = []
    monkeypatch.setattr(
        ocr_agent,
        "_write_workspace_output",
        lambda *args: workspace_writes.append(args) or "/scan.ocr.json",
    )

    result = ocr_agent.run_ocr("read", str(document), str(tmp_path), "json")
    saved = json.loads(workspace_writes[0][2])
    message = ocr_agent.specialist_message(result, "json")

    assert saved["normalized"] == layout
    assert saved["layout_authority"] == "docling"
    assert "surya" not in saved
    assert "glm_ocr" not in saved
    assert result["surya"] == ["text"]
    assert result["glm_ocr"] == ["text"]
    assert "Docling determined the layout and reading order" in message
    assert "agreed on all 1 page" in message
    assert "Paragraph preserved" not in message


def test_graph_declares_top_level_ocr_node_and_edges():
    graph = create_chat_ui()
    assert "ocr" in graph.nodes
    assert "jasper" in graph.nodes
